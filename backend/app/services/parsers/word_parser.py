from __future__ import annotations

import re
from pathlib import Path
from typing import Any

from docx import Document


def parse_word(file_path: Path) -> dict[str, Any]:
    doc = Document(file_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    tables = []
    for table in doc.tables:
        rows = [[cell.text.strip() for cell in row.cells] for row in table.rows]
        tables.append(rows)

    text_content = "\n".join(paragraphs)
    fields = _extract_contract_fields(text_content)

    return {
        "file_type": "word",
        "paragraphs": paragraphs,
        "tables": tables,
        "text_content": text_content,
        "fields": fields,
    }


def _extract_contract_fields(text: str) -> dict[str, Any]:
    fields: dict[str, Any] = {}
    patterns = {
        "contract_no": r"合同编号[:：]?\s*([A-Za-z0-9\-]+)",
        "party_a": r"甲方[:：]?\s*(.+?)(?:\n|乙方)",
        "party_b": r"乙方[:：]?\s*(.+?)(?:\n|签订|合同金额)",
        "contract_amount": r"合同金额[:：]?\s*([\d,\.]+)",
        "tax_rate": r"税率[:：]?\s*([\d\.]+%?)",
        "invoice_type": r"(增值税专用发票|增值税普通发票)",
    }
    for key, pat in patterns.items():
        m = re.search(pat, text, re.S)
        if m:
            val = m.group(1).strip()
            if key == "contract_amount":
                try:
                    val = float(val.replace(",", ""))
                except ValueError:
                    pass
            fields[key] = val
    return fields
