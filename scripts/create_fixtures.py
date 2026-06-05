"""生成完整测试样例集（费用+发票+银行流水+合同）"""
from pathlib import Path

import openpyxl
from docx import Document

ROOT = Path(__file__).resolve().parents[1] / "fixtures"
ROOT.mkdir(exist_ok=True)


def expense_csv():
    p = ROOT / "sample_expense.csv"
    p.write_text(
        """日期,摘要,供应商,金额,发票号,入账日期
2025-01-03,办公用品,A公司,3200,INV001,2025-01-05
2025-02-10,咨询服务费,B公司,128000,,2025-02-15
2025-03-12,差旅费,C公司,8500,INV003,2025-03-12
2025-04-01,业务招待费,D公司,6800,,2025-04-02
2025-12-20,年末咨询,E公司,32000,INV005,2025-01-10
""",
        encoding="utf-8",
    )
    print(f"  {p.name}")


def invoice_csv():
    p = ROOT / "sample_invoice_list.csv"
    p.write_text(
        """发票号码,开票日期,销售方,购买方,价税合计,税额,税率
INV001,2025-01-03,A公司,本公司,3200,184.91,6%
INV003,2025-03-12,C公司,本公司,8500,490.57,6%
INV005,2025-12-18,E公司,本公司,28000,1584.91,6%
INV001,2025-01-03,A公司,本公司,3200,184.91,6%
88888888,2025-02-11,B公司,本公司,106000,6000,6%
""",
        encoding="utf-8",
    )
    print(f"  {p.name}")


def bank_csv():
    p = ROOT / "sample_bank_statement.csv"
    p.write_text(
        """交易日期,对方户名,摘要,借方发生,贷方发生
2025-01-05,A公司,办公用品采购,,3200
2025-02-15,B公司,咨询服务费,,128000
2025-03-12,C公司,差旅报销,,8500
2025-04-10,个人账户,客户款,,50000
2025-06-01,未知,,"",800000,
""",
        encoding="utf-8",
    )
    print(f"  {p.name}")


def contract_docx():
    p = ROOT / "sample_contract.docx"
    doc = Document()
    doc.add_heading("技术服务合同", 0)
    doc.add_paragraph("合同编号：HT-2025-001")
    doc.add_paragraph("甲方：本公司")
    doc.add_paragraph("乙方：B公司")
    doc.add_paragraph("合同金额：500000元")
    doc.add_paragraph("税率：6%")
    doc.add_paragraph("发票类型：增值税专用发票")
    doc.add_paragraph("签订日期：2025-01-01")
    doc.save(p)
    print(f"  {p.name}")


def trial_balance_xlsx():
    p = ROOT / "sample_trial_balance.xlsx"
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "科目余额表"
    ws.append(["科目编码", "科目名称", "借方", "贷方", "余额"])
    ws.append(["1001", "库存现金", 50000, "", 50000])
    ws.append(["2202", "应付账款", "", 128000, 128000])
    ws.append(["6602", "管理费用", 168700, "", 168700])
    wb.save(p)
    print(f"  {p.name}")


def readme():
    p = ROOT / "README.md"
    p.write_text(
        """# 测试样例说明

上传以下文件到同一项目，可验证跨文件比对：

| 文件 | 预期风险 |
|------|---------|
| sample_expense.csv | 大额缺发票、招待费无票 |
| sample_invoice_list.csv | 重复发票 INV001、B公司金额与费用不匹配 |
| sample_bank_statement.csv | 个人账户收款、超大额、摘要空 |
| sample_contract.docx | 合同与发票累计金额不一致 |
| sample_trial_balance.xlsx | 科目余额检查 |

跨期：`sample_expense.csv` 第6行业务日期2024-12-20，入账2025-01-10。
""",
        encoding="utf-8",
    )


if __name__ == "__main__":
    print("生成 fixtures:")
    expense_csv()
    invoice_csv()
    bank_csv()
    contract_docx()
    trial_balance_xlsx()
    readme()
    print("完成")
